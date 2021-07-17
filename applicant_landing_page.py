from tkinter import *
import sqlite3
from tkinter import messagebox


applicant_landing_frame = Tk()
applicant_landing_frame.geometry('500x500')
applicant_landing_frame.title("")

connector=sqlite3.connect('UserRegistrationDetails.db')
with connector:
    cursor=connector.cursor()

cursor.execute("select * from Active_Session")
columns = [item for t in cursor.description for item in t]
columns_final = list(filter(None, columns))
values = [item for t in cursor.fetchall() for item in t]
applicant_dict = dict(zip(columns_final, values))

label_form = Label(applicant_landing_frame, text="Welcome {}!".format(applicant_dict['Name']), width=20, font=("bold", 20))
label_form.place(x=90, y=43)

class applicant_landing:


        def createApplication(self):
            applicant_landing_frame.destroy()
            from application_page import applicant_login_frame
            applicant_login_frame.mainloop()


        def generateOfferLetter(self):
            from docx import Document

            document = Document()

            document.add_heading("Joining Letter", 0)

            p = document.add_paragraph('Congratulations!!! ')

            p1 = document.add_paragraph(
            'We are please to inform you that you have been selected for the applied post.Please find relevent details given below:')

            document.add_paragraph(
                'first item in ordered list', style='List Number'
            )

        # document.add_picture('monty-truth.png', width=Inches(1.25))
            cursor.execute("select * from New_Employee where email='{}'".format(applicant_dict['Email']))
            lst = [item for t in cursor.fetchall() for item in t]
        #lst = ['Kaushal Sanadhya', 'kaushal@gmail.com', '12/12/19', '2345', 'Data Engineer', 'Hyderabad', 'Akshya']
            print(lst)

            table = document.add_table(rows=7, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Email'
            hdr_cells[2].text = 'Date'
            hdr_cells[3].text = 'Salary'
            hdr_cells[4].text = 'Post'
            hdr_cells[5].text = 'Location'
            hdr_cells[6].text = 'Approved_By'

            hdr_cells1 = table.columns[1].cells
            hdr_cells1[0].text = lst[1]
            hdr_cells1[1].text = lst[2]
            hdr_cells1[2].text = lst[3]
            hdr_cells1[3].text = str(lst[4])
            hdr_cells1[4].text = lst[5]
            hdr_cells1[5].text = lst[6]
            hdr_cells1[6].text = lst[7]

        # for i in range(0,len(lst)):
        #    col_cells = table.columns[0].cells
        #    col_cells[i].text =lst[i]

            document.add_page_break()

            document.save('offer_letter.docx')


        def viewApplicationStatus(self):
            temp=cursor.execute("select isShortlisted,InterviewDate,InterviewTime from Applicant_Login where email='{}'".format(applicant_dict['Email']))
            shortlist_status=[item for t in temp.fetchall() for item in t]
            cursor.execute("select count(*) from New_Employee where email='{}'".format(applicant_dict['Email']))
            isHired=int(list(cursor.fetchall()[0])[0])
            if(isHired==1):
                messagebox.showinfo("Congratulations!","Dear Candidate, You have been selected! We will be looking forward for you to join us soon.")
                Button(applicant_landing_frame, text='Download Offer Letter', width=20, bg='brown', fg='white',command=lambda:obj.generateOfferLetter()).place(x=80, y=110)
            elif (shortlist_status[0] == 1 and isHired == 0 and shortlist_status[1] is None):
                messagebox.showinfo("Shortlisted", "Dear Applicant, Congratulations!,You have been shortlisted. We will notify you the date and time of the interview shortly.")
            elif(shortlist_status[0]==1 and isHired==0):
                messagebox.showinfo("Shortlisted","Dear Applicant, You have been shortlisted for interview on {} at {}".format(shortlist_status[1],shortlist_status[2]))
            else:
                messagebox.showinfo("Pending","Dear Applicant, Your application is under review")
obj=applicant_landing()

Button(applicant_landing_frame, text='Create Application', width=20, bg='brown', fg='white',command=lambda:obj.createApplication()).place(x=80, y=80)
Button(applicant_landing_frame, text='View Application Status', width=20, bg='brown', fg='white',command=lambda:obj.viewApplicationStatus()).place(x=80, y=110)

applicant_landing_frame.mainloop()