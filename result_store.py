from tkinter import *
import sqlite3
from tkinter import messagebox

result_store_frame = Tk()
result_store_frame.geometry('500x500')
result_store_frame.title("Rate candidate")

connector=sqlite3.connect('UserRegistrationDetails.db')
with connector:
    cursor=connector.cursor()

applicants=StringVar()
assigned_score=IntVar()

applicant_temp = []
cursor.execute("select FirstName from Applicant_Login where isShortlisted=1")
applicant_temp=[item for t in cursor.fetchall() for item in t]


class Result_Store:

    if(len(applicant_temp)!=0):
        applicants.set(applicant_temp[0])
        applicant_options = OptionMenu(result_store_frame, applicants, *applicant_temp)
        applicant_options.place(x=220, y=85)


        def generateInterviewReport(self,score,email):
            from docx import Document
            from docx.shared import Inches

            document = Document()

            document.add_heading("Interview Report", 0)

            p = document.add_paragraph('Congratulations!!! ')

            p1 = document.add_paragraph(
                'Please find the interview report of the candidate below:')

            document.add_paragraph(
                'first item in ordered list', style='List Number'
            )

            # document.add_picture('monty-truth.png', width=Inches(1.25))
            cursor.execute("select FirstName,LastName,VacancyApplied,HR_Assigned,InterviewDate,InterviewTime,InterviewScore from Applicant_Login where email='{}'".format(email))
            lst = [item for t in cursor.fetchall() for item in t]
            # lst = ['Kaushal Sanadhya', 'kaushal@gmail.com', '12/12/19', '2345', 'Data Engineer', 'Hyderabad', 'Akshya']
            print(lst)

            table = document.add_table(rows=7, cols=2)
            hdr_cells = table.columns[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Role'
            hdr_cells[2].text = 'HR Representative'
            hdr_cells[3].text = 'Interview Date'
            hdr_cells[4].text = 'Interview Time'
            hdr_cells[5].text = 'Interview Score'

            hdr_cells1 = table.columns[1].cells
            hdr_cells1[0].text = lst[0]+lst[1]
            hdr_cells1[1].text = lst[2]
            hdr_cells1[2].text = lst[3]
            hdr_cells1[3].text = str(lst[4])
            hdr_cells1[4].text = str(lst[5])
            hdr_cells1[5].text = str(lst[6])


            # for i in range(0,len(lst)):
            #    col_cells = table.columns[0].cells
            #    col_cells[i].text =lst[i]

            document.add_page_break()

            document.save('interview_report.docx')


        def storeResult(self,email):
            score=assigned_score.get()
            fields_validity=1
            if(score==''):
                messagebox.showinfo('Error', 'All fields are required')
                field_validity = 0
            else:
                cursor.execute("update Applicant_login set InterviewScore={} where Email='{}'".format(score,email))
                messagebox.showinfo('Success',"Applicant Rating Successful")
                connector.commit()
                Button(result_store_frame, text='Generate Interview Report', width=20, bg='brown', fg='white',
                       command=lambda:obj. generateInterviewReport(score,email)).place(x=260, y=320)
                #result_store_frame.destroy()



        def rateCandidate(*args):
            cursor.execute(
                "select FirstName, LastName, Email, VacancyApplied from Applicant_Login where FirstName='{}'".format(
                    applicants.get()))
            columns = [item for t in cursor.description for item in t]
            columns_final = list(filter(None, columns))
            values = [item for t in cursor.fetchall() for item in t]
            applicant_dict = dict(zip(columns_final, values))

            label_name = Label(result_store_frame, text="Name", width=20, font=("bold", 10))
            label_name.place(x=80, y=140)

            name_text = Entry(result_store_frame)
            name_text.insert('end', applicant_dict['FirstName'] + applicant_dict['LastName'])
            name_text.config(state=DISABLED)
            name_text.place(x=240, y=140)

            label_vacancy = Label(result_store_frame, text="Vacancy applied", width=20, font=("bold", 10))
            label_vacancy.place(x=80, y=165)

            vacancy_text = Entry(result_store_frame)
            vacancy_text.insert('end', applicant_dict['VacancyApplied'])
            vacancy_text.config(state=DISABLED)
            vacancy_text.place(x=240, y=165)

            rating=[1,2,3,4,5]

            label_score = Label(result_store_frame, text="Overall Rating", width=20, font=("bold", 10))
            label_score.place(x=80, y=195)

            score_options = OptionMenu(result_store_frame, assigned_score, *rating)
            score_options.place(x=240, y=190)

            Button(result_store_frame, text='Store Result', width=20, bg='brown', fg='white',
                   command=lambda: obj.storeResult(applicant_dict['Email'])).place(x=80, y=320)

        applicants.trace('w', rateCandidate)

    else:
        label_none = Label(result_store_frame, text="No Shortlisted Candidates", width=20, font=("bold", 12))
        label_none.place(x=150, y=140)
obj=Result_Store()
result_store_frame.mainloop()