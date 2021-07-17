from tkinter import *
import sqlite3
from tkinter import messagebox

initiate_interview_frame = Tk()
initiate_interview_frame.geometry('500x500')
initiate_interview_frame.title("Initiate Interview")

connector=sqlite3.connect('UserRegistrationDetails.db')
with connector:
    cursor=connector.cursor()

applicants=StringVar()

label_form = Label(initiate_interview_frame, text="Initiate Interview", width=20, font=("bold", 20))
label_form.place(x=90, y=43)

applicant_temp = []
cursor.execute("select FirstName from Applicant_Login where isEligible=1")
applicant_temp=[item for t in cursor.fetchall() for item in t]
#print(applicant_temp)

class Initiate_Interview:

    if(len(applicant_temp)!=0):
        applicants.set(applicant_temp[0])
        applicant_options = OptionMenu(initiate_interview_frame, applicants, *applicant_temp)
        applicant_options.place(x=220, y=85)


        def generateShortlist(self,email):
            #cursor.execute('select ')
            from docx import Document
            from docx.shared import Inches

            document = Document()

            document.add_heading("List of Candidates Short Listed", 0)

            p1 = document.add_paragraph(
                'Please find the list of shortlisted candidates below:')

            # document.add_picture('monty-truth.png', width=Inches(1.25))
            cursor.execute("select FirstName from Applicant_Login where isShortlisted=1")
            lst = [item for t in cursor.fetchall() for item in t]
            # lst = ['Kaushal Sanadhya', 'kaushal@gmail.com', '12/12/19', '2345', 'Data Engineer', 'Hyderabad', 'Akshya']

            table = document.add_table(rows=len(lst), cols=1)
            hdr_cells = table.columns[0].cells

            for index,item in enumerate(lst):
                hdr_cells[index].text = lst[index]

            # for i in range(0,len(lst)):
            #    col_cells = table.columns[0].cells
            #    col_cells[i].text =lst[i]

            document.add_page_break()

            document.save('shortlisted_candidates.docx')


        def shortlistCandidate(self,val,email):
            if val==0:
                cursor.execute("delete from Applicant_Login where email='{}'".format(email))
                messagebox.showinfo('Success', 'Candidate Rejected')
            else:
                cursor.execute("update Applicant_Login set isShortlisted=1 where email='{}'".format(email))
                messagebox.showinfo('Success', 'Candidate Shortlisted')
            connector.commit()
            initiate_interview_frame.destroy()

        def populate_user_details(*args):
            cursor.execute("select FirstName, LastName, Email, VacancyApplied,Competencies,WorkExperience,Salary,Tools_Research_Papers from Applicant_Login where FirstName='{}'".format(applicants.get()))
            columns = [item for t in cursor.description for item in t]
            columns_final = list(filter(None, columns))
            values = [item for t in cursor.fetchall() for item in t]
            applicant_dict = dict(zip(columns_final, values))

            label_name = Label(initiate_interview_frame, text="Name", width=20, font=("bold", 10))
            label_name.place(x=80, y=140)

            name_text = Entry(initiate_interview_frame)
            name_text.insert('end', applicant_dict['FirstName']+applicant_dict['LastName'])
            name_text.config(state=DISABLED)
            name_text.place(x=240, y=140)

            label_vacancy = Label(initiate_interview_frame, text="Vacancy applied", width=20, font=("bold", 10))
            label_vacancy.place(x=80, y=165)

            vacancy_text = Entry(initiate_interview_frame)
            vacancy_text.insert('end', applicant_dict['VacancyApplied'])
            vacancy_text.config(state=DISABLED)
            vacancy_text.place(x=240, y=165)

            label_skill = Label(initiate_interview_frame, text="Competencies", width=20, font=("bold", 10))
            label_skill.place(x=80, y=190)

            skill_text = Entry(initiate_interview_frame)
            skill_text.insert('end', applicant_dict['Competencies'])
            skill_text.config(state=DISABLED)
            skill_text.place(x=240, y=190)

            label_exp = Label(initiate_interview_frame, text="Work Experience", width=20, font=("bold", 10))
            label_exp.place(x=80, y=215)

            exp_text = Entry(initiate_interview_frame)
            exp_text.insert('end', applicant_dict['WorkExperience'])
            exp_text.config(state=DISABLED)
            exp_text.place(x=240, y=215)

            label_sal = Label(initiate_interview_frame, text="Salary", width=20, font=("bold", 10))
            label_sal.place(x=80, y=240)

            sal_text = Entry(initiate_interview_frame)
            sal_text.insert('end', applicant_dict['Salary'])
            sal_text.config(state=DISABLED)
            sal_text.place(x=240, y=240)

            label_misc = Label(initiate_interview_frame, text="Tools/Research_Papers", width=20, font=("bold", 10))
            label_misc.place(x=80, y=265)

            misc_text = Entry(initiate_interview_frame)
            misc_text.insert('end', applicant_dict['Tools_Research_Papers'])
            misc_text.config(state=DISABLED)
            misc_text.place(x=240, y=265)

            Button(initiate_interview_frame, text='Shortlist Candidate', width=20, bg='brown', fg='white',command=lambda:obj.shortlistCandidate(1,applicant_dict['Email'])).place(x=80, y=320)
            Button(initiate_interview_frame, text='Reject Candidate', width=20, bg='brown', fg='white',command=lambda:obj.shortlistCandidate(0,applicant_dict['Email'])).place(x=250, y=320)
            Button(initiate_interview_frame, text='Shortlisted Candidates Report', width=20, bg='brown', fg='white',
                   command=lambda: obj.generateShortlist(applicant_dict['Email'])).place(x=160, y=350)
        applicants.trace('w', populate_user_details)
    else:
        label_none = Label(initiate_interview_frame, text="No Eligible Candidates", width=20, font=("bold", 12))
        label_none.place(x=150, y=140)
obj=Initiate_Interview()
initiate_interview_frame.mainloop()
